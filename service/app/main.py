"""Servicio FastAPI — ingesta, extraccion, segmentacion y clasificacion (Art IV, pasos 1-5).

El pipeline se detiene despues de la clasificacion (paso 5): sin score de confianza,
sin cola de revision, sin publicacion (spec-mvp-demo.md).

POST /documentos responde 201 apenas persiste el documento y corre el pipeline en una
tarea en segundo plano (BackgroundTasks), NO dentro de la request: la clasificacion hace
una llamada al modelo por clausula, y en un documento real eso excede el timeout del
ingress de Container Apps — la request moria y el navegador lo reportaba como un error de
CORS (la respuesta de error del proxy no lleva cabeceras CORS). El avance se sigue por la
columna `estado` (pendiente -> extraido -> segmentado -> clasificado | error), que ya
existia para eso. Sigue sin haber cola de tareas (Azure Service Bus, Art V) — sigue siendo
la simplificacion de demo señalada explicitamente.
"""
import uuid
from pathlib import Path
from typing import Optional

import httpx
from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware

from app import storage
from app.auth import require_role
from app.classification import build_system_prompt, check_legal_compliance, classify_clause
from app.config import WEB_ORIGIN
from app.db import get_conn
from app.extraction import ExtractionError, extract_text
from app.segmentation import segment_clauses

app = FastAPI(title="Comparador de Documentos Legales — demo Venezuela")

# La app de React (web/) llama a este servicio directo desde el navegador — sin esto el
# navegador bloquea la carga/lista de documentos con "No 'Access-Control-Allow-Origin'".
app.add_middleware(
    CORSMiddleware,
    allow_origins=[WEB_ORIGIN],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------------------------------------------------------------------
# Tenants (spec-mvp-demo.md #3: alta de tenant = alta de empresa, pais fijo Venezuela)
# ---------------------------------------------------------------------------


@app.post("/tenants", status_code=201)
def crear_tenant(nombre_empresa: str = Form(...)):
    with get_conn() as conn, conn.cursor() as cur:
        # Pais fijo Venezuela para este MVP (Art I.3): sin selector, se resuelve el unico
        # pais activo de la tabla paises en vez de pedirlo en el form.
        cur.execute(
            """
            INSERT INTO tenants (nombre_empresa, pais_id)
            SELECT %s, id FROM paises WHERE codigo = 'VE'
            RETURNING id, nombre_empresa, pais_id, plan_licencia, created_at
            """,
            (nombre_empresa,),
        )
        tenant = cur.fetchone()
        conn.commit()
    return tenant


@app.get("/tenants")
def listar_tenants():
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            """
            SELECT t.id, t.nombre_empresa, t.plan_licencia, t.created_at,
                   p.codigo AS pais_codigo, p.nombre AS pais_nombre
            FROM tenants t
            JOIN paises p ON p.id = t.pais_id
            ORDER BY t.created_at
            """
        )
        return cur.fetchall()


# ---------------------------------------------------------------------------
# Taxonomia (Art II): de solo lectura -- se usa para poblar el selector de "corregir"
# en la cola de revision (Bloque D). La carga real vive en db/seed_taxonomia.py.
# ---------------------------------------------------------------------------


@app.get("/taxonomia")
def listar_taxonomia():
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            """
            SELECT t.id, t.nombre, c.id AS categoria_id, c.nombre AS categoria_nombre
            FROM taxonomia_titulos t
            JOIN taxonomia_categorias c ON c.id = t.categoria_id
            ORDER BY c.nombre, t.nombre
            """
        )
        return cur.fetchall()


# ---------------------------------------------------------------------------
# Catalogos globales de segmentacion de empresas (Art II.5, spec-empresas-comparacion.md
# Bloque A) -- compartidos por todos los tenants, de solo lectura: su administracion queda
# para el rol de Plataforma (Fase 5), ningun rol de tenant puede editarlos en esta fase.
# ---------------------------------------------------------------------------


@app.get("/catalogos")
def listar_catalogos():
    with get_conn() as conn, conn.cursor() as cur:
        catalogos = {}
        for tabla in ["sectores", "tipos_empresa", "categorias_sector", "actividades_empresa"]:
            cur.execute(f"SELECT id, nombre, descripcion FROM {tabla} ORDER BY nombre")
            catalogos[tabla] = cur.fetchall()

        cur.execute("SELECT id, nombre FROM estados ORDER BY nombre")
        catalogos["estados"] = cur.fetchall()

    return catalogos


@app.get("/catalogos/localidades")
def listar_localidades(estado_id: int):
    # Filtrado por estado (no todas las 409 de una vez): mismo patron cascada que el
    # legado (sector -> tipo -> ... -> estado -> localidad, comparador.php).
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            "SELECT id, nombre FROM localidades WHERE estado_id = %s ORDER BY nombre",
            (estado_id,),
        )
        return cur.fetchall()


# ---------------------------------------------------------------------------
# Empresas (Art III de la constitucion, spec-empresas-comparacion.md Bloque B): el
# catalogo propio de un tenant (operador) -- la esencia del producto original, adaptada de
# empresas.php del legado. Los campos de segmentacion (sector/tipo/categoria/actividad/
# estado/localidad) son opcionales a proposito: cargar una empresa no deberia bloquearse
# si todavia no se conocen todos esos datos.
# ---------------------------------------------------------------------------

_EMPRESA_CAMPOS_SELECT = """
    e.id, e.nombre, e.rif, e.contacto_nombre, e.contacto_email, e.created_at,
    e.sector_id, s.nombre AS sector_nombre,
    e.tipo_id, t.nombre AS tipo_nombre,
    e.categoria_id, c.nombre AS categoria_nombre,
    e.actividad_id, a.nombre AS actividad_nombre,
    e.estado_id, es.nombre AS estado_nombre,
    e.localidad_id, l.nombre AS localidad_nombre
"""
_EMPRESA_JOINS = """
    FROM empresas e
    LEFT JOIN sectores s ON s.id = e.sector_id
    LEFT JOIN tipos_empresa t ON t.id = e.tipo_id
    LEFT JOIN categorias_sector c ON c.id = e.categoria_id
    LEFT JOIN actividades_empresa a ON a.id = e.actividad_id
    LEFT JOIN estados es ON es.id = e.estado_id
    LEFT JOIN localidades l ON l.id = e.localidad_id
"""


@app.post("/empresas", status_code=201)
def crear_empresa(
    request: Request,
    nombre: str = Form(...),
    rif: Optional[str] = Form(None),
    sector_id: Optional[int] = Form(None),
    tipo_id: Optional[int] = Form(None),
    categoria_id: Optional[int] = Form(None),
    actividad_id: Optional[int] = Form(None),
    estado_id: Optional[int] = Form(None),
    localidad_id: Optional[int] = Form(None),
    contacto_nombre: Optional[str] = Form(None),
    contacto_email: Optional[str] = Form(None),
):
    # auth-spec.md §5 (extendido en spec-empresas-comparacion.md §3): gestionar el catalogo
    # de empresas es AdminTenant/Editor, igual que cargar documentos. tenant_id sale del
    # claim del JWT (Art VI.2), nunca de un form field.
    claims = require_role(request, "AdminTenant", "Editor")
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            """
            INSERT INTO empresas
                (tenant_id, nombre, rif, sector_id, tipo_id, categoria_id, actividad_id,
                 estado_id, localidad_id, contacto_nombre, contacto_email)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id
            """,
            (
                claims.tenant_id, nombre, rif, sector_id, tipo_id, categoria_id,
                actividad_id, estado_id, localidad_id, contacto_nombre, contacto_email,
            ),
        )
        empresa_id = cur.fetchone()["id"]
        conn.commit()
    return _obtener_empresa(empresa_id, claims.tenant_id)


def _obtener_empresa(empresa_id, tenant_id) -> dict:
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            f"SELECT {_EMPRESA_CAMPOS_SELECT} {_EMPRESA_JOINS} WHERE e.id = %s AND e.tenant_id = %s",
            (empresa_id, tenant_id),
        )
        empresa = cur.fetchone()
        if empresa is None:
            raise HTTPException(404, "empresa no encontrada")
        return empresa


@app.get("/empresas")
def listar_empresas(request: Request):
    claims = require_role(request, "AdminTenant", "Editor")
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            f"SELECT {_EMPRESA_CAMPOS_SELECT} {_EMPRESA_JOINS} WHERE e.tenant_id = %s ORDER BY e.nombre",
            (claims.tenant_id,),
        )
        return cur.fetchall()


@app.get("/empresas/{empresa_id}")
def obtener_empresa(request: Request, empresa_id: uuid.UUID):
    claims = require_role(request, "AdminTenant", "Editor")
    return _obtener_empresa(empresa_id, claims.tenant_id)


@app.put("/empresas/{empresa_id}")
def editar_empresa(
    request: Request,
    empresa_id: uuid.UUID,
    nombre: str = Form(...),
    rif: Optional[str] = Form(None),
    sector_id: Optional[int] = Form(None),
    tipo_id: Optional[int] = Form(None),
    categoria_id: Optional[int] = Form(None),
    actividad_id: Optional[int] = Form(None),
    estado_id: Optional[int] = Form(None),
    localidad_id: Optional[int] = Form(None),
    contacto_nombre: Optional[str] = Form(None),
    contacto_email: Optional[str] = Form(None),
):
    claims = require_role(request, "AdminTenant", "Editor")
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            """
            UPDATE empresas SET
                nombre = %s, rif = %s, sector_id = %s, tipo_id = %s, categoria_id = %s,
                actividad_id = %s, estado_id = %s, localidad_id = %s,
                contacto_nombre = %s, contacto_email = %s
            WHERE id = %s AND tenant_id = %s
            """,
            (
                nombre, rif, sector_id, tipo_id, categoria_id, actividad_id, estado_id,
                localidad_id, contacto_nombre, contacto_email, empresa_id, claims.tenant_id,
            ),
        )
        if cur.rowcount == 0:
            raise HTTPException(404, "empresa no encontrada")
        conn.commit()
    return _obtener_empresa(empresa_id, claims.tenant_id)


# ---------------------------------------------------------------------------
# Documentos: ingesta + pipeline (Art IV, pasos 1-5)
# ---------------------------------------------------------------------------


def _validar_url_publica(url: str) -> None:
    """Art IV.2 / VI.1: si se declara publico via URL, debe validarse que responda
    sin autenticacion antes de aceptarla como publica."""
    try:
        resp = httpx.get(url, timeout=15.0, follow_redirects=True)
    except httpx.HTTPError as exc:
        raise HTTPException(
            422,
            f"No se pudo marcar el documento como publico: la URL no respondio ({exc}). "
            "Corrige la URL o vuelve a intentar sin marcarlo como publico.",
        ) from exc
    if resp.status_code != 200:
        raise HTTPException(
            422,
            f"No se pudo marcar el documento como publico: la URL respondio {resp.status_code} "
            "(se esperaba 200 sin autenticacion). Corrige la URL o vuelve a intentar sin marcarlo como publico.",
        )


def _descargar_url(url: str) -> tuple[bytes, str]:
    try:
        resp = httpx.get(url, timeout=30.0, follow_redirects=True)
        resp.raise_for_status()
    except httpx.HTTPError as exc:
        raise HTTPException(422, f"No se pudo descargar el documento desde la URL: {exc}") from exc

    nombre = Path(url.split("?")[0]).name or "documento"
    if "." not in nombre:
        content_type = resp.headers.get("content-type", "")
        if "pdf" in content_type:
            nombre += ".pdf"
        elif "word" in content_type or "officedocument.wordprocessingml" in content_type:
            nombre += ".docx"
    return resp.content, nombre


@app.post("/documentos", status_code=201)
def crear_documento(
    request: Request,
    background: BackgroundTasks,
    empresa_id: uuid.UUID = Form(...),
    origen: str = Form(...),
    url_origen: Optional[str] = Form(None),
    es_publico: bool = Form(False),
    archivo: Optional[UploadFile] = File(None),
):
    # auth-spec.md §5: "Cargar documento (ingesta)" = AdminTenant/Editor. tenant_id sale
    # del claim del JWT (Art VI.2), nunca de un form field — evita cross-tenant manipulando
    # el payload.
    claims = require_role(request, "AdminTenant", "Editor")
    tenant_id = claims.tenant_id

    if origen not in ("archivo", "url"):
        raise HTTPException(422, "origen debe ser 'archivo' o 'url'")
    if origen == "url" and not url_origen:
        raise HTTPException(422, "url_origen es requerido cuando origen='url'")
    if origen == "archivo" and archivo is None:
        raise HTTPException(422, "archivo es requerido cuando origen='archivo'")
    if es_publico and origen != "url":
        raise HTTPException(422, "es_publico solo aplica a documentos ingresados por URL")

    with get_conn() as conn, conn.cursor() as cur:
        cur.execute("SELECT id FROM tenants WHERE id = %s", (tenant_id,))
        if cur.fetchone() is None:
            raise HTTPException(404, f"tenant_id {tenant_id} no existe")

        # Bloque C: empresa_id debe pertenecer al catalogo de ESTE tenant (Art VI.2) --
        # nunca confiar en que el id que manda el cliente sea de su propio tenant.
        cur.execute("SELECT id FROM empresas WHERE id = %s AND tenant_id = %s", (empresa_id, tenant_id))
        if cur.fetchone() is None:
            raise HTTPException(404, f"empresa_id {empresa_id} no existe en tu catálogo")

    if es_publico:
        _validar_url_publica(url_origen)

    # contenido/extension se mantienen en memoria para el pipeline (extraction.py) — nunca
    # se vuelve a leer el original desde su ubicacion persistida (storage.guardar) para
    # procesarlo, asi que scale-to-zero o una replica distinta entre requests no rompe nada.
    ruta_archivo = None
    contenido = None
    extension = None
    if origen == "archivo":
        contenido = archivo.file.read()
        extension = Path(archivo.filename).suffix
        ruta_archivo = storage.guardar(tenant_id, archivo.filename, contenido)

    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            """
            INSERT INTO documentos (tenant_id, empresa_id, origen, url_origen, ruta_archivo, es_publico, estado)
            VALUES (%s, %s, %s, %s, %s, %s, 'pendiente')
            RETURNING id
            """,
            (tenant_id, empresa_id, origen, url_origen, ruta_archivo, es_publico),
        )
        documento_id = cur.fetchone()["id"]
        conn.commit()

    if origen == "url":
        contenido, nombre = _descargar_url(url_origen)
        extension = Path(nombre).suffix
        ruta_archivo = storage.guardar(tenant_id, nombre, contenido)
        with get_conn() as conn, conn.cursor() as cur:
            cur.execute("UPDATE documentos SET ruta_archivo = %s WHERE id = %s", (ruta_archivo, documento_id))
            conn.commit()

    # En segundo plano, despues de responder: el documento queda en 'pendiente' y el
    # cliente sigue el avance por la columna `estado` (la lista se refresca sola).
    background.add_task(_procesar_pipeline, documento_id, tenant_id, contenido, extension)

    return _obtener_documento(documento_id, tenant_id)


def _marcar_error(documento_id: int, mensaje: str) -> None:
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            "UPDATE documentos SET estado = 'error', estado_detalle = %s WHERE id = %s",
            (mensaje, documento_id),
        )
        conn.commit()


def _articulos_relacionados_a_titulo(cur, titulo_id: int, tenant_id: uuid.UUID) -> list[dict]:
    """Articulos de ley vinculados a un titulo (Fase 4, spec-marco-legal.md), filtrados por
    el pais del tenant (Art VI.2 -- ninguna razon para cruzar cumplimiento contra la ley de
    un pais distinto al del tenant, aunque hoy solo exista corpus de Venezuela)."""
    cur.execute(
        """
        SELECT al.nro_articulo, al.titulo_articulo, al.texto_completo
        FROM titulo_articulo_ley tal
        JOIN articulos_ley al ON al.id = tal.articulo_ley_id
        JOIN leyes l ON l.id = al.ley_id
        JOIN tenants t ON t.pais_id = l.pais_id
        WHERE tal.titulo_id = %s AND t.id = %s
        ORDER BY al.nro_articulo
        """,
        (titulo_id, tenant_id),
    )
    return cur.fetchall()


def _procesar_pipeline(documento_id: int, tenant_id: uuid.UUID, contenido: bytes, extension: str) -> None:
    try:
        texto = extract_text(contenido, extension)
    except ExtractionError as exc:
        _marcar_error(documento_id, str(exc))
        return

    with get_conn() as conn, conn.cursor() as cur:
        cur.execute("UPDATE documentos SET estado = 'extraido' WHERE id = %s", (documento_id,))
        conn.commit()

    clausulas_texto = segment_clauses(texto)
    if not clausulas_texto:
        _marcar_error(documento_id, "No se pudo segmentar ninguna clausula del texto extraido.")
        return

    with get_conn() as conn, conn.cursor() as cur:
        cur.execute("UPDATE documentos SET estado = 'segmentado' WHERE id = %s", (documento_id,))
        conn.commit()

        cur.execute(
            """
            SELECT t.id, t.nombre, t.descripcion, c.id AS categoria_id, c.nombre AS categoria_nombre
            FROM taxonomia_titulos t
            JOIN taxonomia_categorias c ON c.id = t.categoria_id
            ORDER BY c.id, t.id
            """
        )
        titulos = cur.fetchall()

    system_prompt = build_system_prompt(titulos)
    fallos = 0

    with get_conn() as conn, conn.cursor() as cur:
        for orden, texto_clausula in enumerate(clausulas_texto, start=1):
            titulo_id = None
            categoria_id = None
            confianza = None
            try:
                resultado = classify_clause(texto_clausula, titulos, system_prompt)
                titulo_id = resultado["titulo_id"]
                categoria_id = resultado["categoria_id"]
                confianza = resultado["confianza"]
            except Exception as exc:  # nunca abortar todo el documento por una clausula
                print(f"[clasificacion] documento {documento_id} orden {orden}: {exc}")
                fallos += 1

            # Art IV.5 bis (spec-marco-legal.md): solo se llama al modelo si el titulo ya
            # asignado tiene articulos de ley vinculados -- si no, 'no_aplica' sin gastar
            # una llamada extra. Nunca bloquea el pipeline: un fallo acá deja la señal en
            # NULL, igual que un fallo de clasificacion deja titulo_id en NULL.
            cumplimiento_legal = None
            cumplimiento_justificacion = None
            if titulo_id is not None:
                articulos_relacionados = _articulos_relacionados_a_titulo(cur, titulo_id, tenant_id)
                if not articulos_relacionados:
                    cumplimiento_legal = "no_aplica"
                else:
                    try:
                        resultado_legal = check_legal_compliance(texto_clausula, articulos_relacionados)
                        cumplimiento_legal = resultado_legal["cumplimiento"]
                        cumplimiento_justificacion = resultado_legal["justificacion"]
                    except Exception as exc:
                        print(f"[cumplimiento] documento {documento_id} orden {orden}: {exc}")

            cur.execute(
                """
                INSERT INTO clausulas
                    (documento_id, tenant_id, texto, titulo_id, categoria_id, orden, confianza,
                     cumplimiento_legal, cumplimiento_justificacion)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                """,
                (
                    documento_id, tenant_id, texto_clausula, titulo_id, categoria_id, orden, confianza,
                    cumplimiento_legal, cumplimiento_justificacion,
                ),
            )
        conn.commit()

    detalle = None
    if fallos:
        detalle = f"{fallos} de {len(clausulas_texto)} clausulas no pudieron clasificarse automaticamente."

    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            "UPDATE documentos SET estado = 'clasificado', estado_detalle = %s WHERE id = %s",
            (detalle, documento_id),
        )
        conn.commit()


def _obtener_documento(documento_id: int, tenant_id: uuid.UUID) -> dict:
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            "SELECT * FROM documentos WHERE id = %s AND tenant_id = %s",
            (documento_id, tenant_id),
        )
        documento = cur.fetchone()
        if documento is None:
            raise HTTPException(404, "documento no encontrado")

        cur.execute(
            """
            SELECT cl.id, cl.texto, cl.orden, t.nombre AS titulo_nombre, c.nombre AS categoria_nombre,
                   cl.confianza, cl.estado_revision, cl.cumplimiento_legal, cl.cumplimiento_justificacion
            FROM clausulas cl
            LEFT JOIN taxonomia_titulos t ON t.id = cl.titulo_id
            LEFT JOIN taxonomia_categorias c ON c.id = cl.categoria_id
            WHERE cl.documento_id = %s AND cl.tenant_id = %s
            ORDER BY cl.orden
            """,
            (documento_id, tenant_id),
        )
        documento["clausulas"] = cur.fetchall()
    return documento


@app.get("/documentos/{documento_id}")
def obtener_documento(request: Request, documento_id: int):
    # Art VI.2: tenant_id sale SIEMPRE del claim del JWT, nunca de un query param — antes
    # este endpoint tomaba tenant_id directo de la URL, sin ninguna autenticacion.
    claims = require_role(request, "AdminTenant", "Revisor", "Editor", "Visualizador")
    return _obtener_documento(documento_id, claims.tenant_id)


def _listar_documentos(tenant_id: uuid.UUID) -> list[dict]:
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            """
            SELECT d.id, d.origen, d.url_origen, d.es_publico, d.estado, d.estado_detalle,
                   d.created_at, d.empresa_id, e.nombre AS empresa_nombre
            FROM documentos d
            JOIN empresas e ON e.id = d.empresa_id
            WHERE d.tenant_id = %s
            ORDER BY d.id DESC
            """,
            (tenant_id,),
        )
        return cur.fetchall()


@app.get("/documentos")
def listar_documentos(request: Request):
    claims = require_role(request, "AdminTenant", "Revisor", "Editor", "Visualizador")
    return _listar_documentos(claims.tenant_id)


# ---------------------------------------------------------------------------
# Cola de revision (Art IV.7-8 de la constitucion, no negociable desde v1.0.0 pero nunca
# construida hasta ahora / spec-empresas-comparacion.md Bloque D). Ninguna clausula queda
# visible para comparacion (Art IV.9) sin pasar por aca primero.
# ---------------------------------------------------------------------------


@app.get("/revision")
def listar_cola_revision(request: Request):
    claims = require_role(request, "AdminTenant", "Revisor")
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            """
            SELECT cl.id, cl.texto, cl.orden, cl.confianza,
                   cl.documento_id, d.empresa_id, e.nombre AS empresa_nombre,
                   cl.titulo_id, t.nombre AS titulo_nombre, cl.categoria_id, c.nombre AS categoria_nombre,
                   cl.cumplimiento_legal, cl.cumplimiento_justificacion
            FROM clausulas cl
            JOIN documentos d ON d.id = cl.documento_id
            JOIN empresas e ON e.id = d.empresa_id
            LEFT JOIN taxonomia_titulos t ON t.id = cl.titulo_id
            LEFT JOIN taxonomia_categorias c ON c.id = cl.categoria_id
            WHERE cl.tenant_id = %s AND cl.estado_revision = 'pendiente'
            ORDER BY
                -- confianza baja (o sin confianza, ej. fallo de clasificacion) primero:
                -- Art IV.7, prioriza la cola por lo que mas necesita ojo humano.
                CASE cl.confianza WHEN 'bajo' THEN 0 WHEN 'medio' THEN 1 WHEN 'alto' THEN 2 ELSE -1 END,
                cl.documento_id, cl.orden
            """,
            (claims.tenant_id,),
        )
        return cur.fetchall()


@app.post("/revision/{clausula_id}/aprobar")
def aprobar_clausula(
    request: Request,
    clausula_id: int,
    titulo_id: Optional[int] = Form(None),
):
    # "Corregir" (auth-spec.md §5) se resuelve aca mismo: si el Revisor manda un titulo_id
    # distinto al sugerido, se actualiza como parte de la aprobacion -- no es una accion
    # separada, es el mismo gesto de "reviso y confirmo (con o sin ajuste)".
    claims = require_role(request, "AdminTenant", "Revisor")
    with get_conn() as conn, conn.cursor() as cur:
        if titulo_id is not None:
            cur.execute("SELECT categoria_id FROM taxonomia_titulos WHERE id = %s", (titulo_id,))
            fila = cur.fetchone()
            if fila is None:
                raise HTTPException(422, f"titulo_id {titulo_id} no existe en la taxonomía")
            cur.execute(
                "UPDATE clausulas SET titulo_id = %s, categoria_id = %s WHERE id = %s AND tenant_id = %s",
                (titulo_id, fila["categoria_id"], clausula_id, claims.tenant_id),
            )

        cur.execute(
            """
            UPDATE clausulas
            SET estado_revision = 'aprobado', revisado_por = %s, revisado_at = now()
            WHERE id = %s AND tenant_id = %s
            RETURNING id
            """,
            (claims.user_id, clausula_id, claims.tenant_id),
        )
        if cur.fetchone() is None:
            raise HTTPException(404, "clausula no encontrada")
        conn.commit()
    return {"id": clausula_id, "estado_revision": "aprobado"}


@app.post("/revision/{clausula_id}/rechazar")
def rechazar_clausula(request: Request, clausula_id: int):
    claims = require_role(request, "AdminTenant", "Revisor")
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            """
            UPDATE clausulas
            SET estado_revision = 'rechazado', revisado_por = %s, revisado_at = now()
            WHERE id = %s AND tenant_id = %s
            RETURNING id
            """,
            (claims.user_id, clausula_id, claims.tenant_id),
        )
        if cur.fetchone() is None:
            raise HTTPException(404, "clausula no encontrada")
        conn.commit()
    return {"id": clausula_id, "estado_revision": "rechazado"}


# ---------------------------------------------------------------------------
# Comparador (Art III "Reporte de comparación", Art IV.9 / spec-empresas-comparacion.md
# Bloque E) -- el nucleo del producto original (comparador.php del legado), ahora
# intra-tenant: compara SOLO dentro del catalogo de empresas de un mismo tenant, y SOLO
# clausulas ya aprobadas por revision humana (Art IV.8/IV.9 -- nunca se publica sin eso).
# ---------------------------------------------------------------------------


@app.get("/comparador/titulos")
def listar_titulos_con_aprobadas(request: Request):
    """Titulos que tienen al menos una clausula aprobada en este tenant -- para poblar el
    selector principal del comparador sin ofrecer titulos vacios."""
    claims = require_role(request, "AdminTenant", "Revisor", "Editor", "Visualizador")
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            """
            SELECT DISTINCT t.id, t.nombre, c.nombre AS categoria_nombre
            FROM clausulas cl
            JOIN taxonomia_titulos t ON t.id = cl.titulo_id
            JOIN taxonomia_categorias c ON c.id = t.categoria_id
            WHERE cl.tenant_id = %s AND cl.estado_revision = 'aprobado'
            ORDER BY c.nombre, t.nombre
            """,
            (claims.tenant_id,),
        )
        return cur.fetchall()


@app.get("/comparador")
def comparar(
    request: Request,
    titulo_id: int,
    sector_id: Optional[int] = None,
    tipo_id: Optional[int] = None,
    categoria_id: Optional[int] = None,
    actividad_id: Optional[int] = None,
    estado_id: Optional[int] = None,
):
    claims = require_role(request, "AdminTenant", "Revisor", "Editor", "Visualizador")
    filtros = ["cl.tenant_id = %s", "cl.titulo_id = %s", "cl.estado_revision = 'aprobado'"]
    params: list = [claims.tenant_id, titulo_id]
    for campo, valor in [
        ("e.sector_id", sector_id), ("e.tipo_id", tipo_id), ("e.categoria_id", categoria_id),
        ("e.actividad_id", actividad_id), ("e.estado_id", estado_id),
    ]:
        if valor is not None:
            filtros.append(f"{campo} = %s")
            params.append(valor)

    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            f"""
            SELECT e.id AS empresa_id, e.nombre AS empresa_nombre,
                   cl.id AS clausula_id, cl.texto, cl.orden, d.id AS documento_id
            FROM clausulas cl
            JOIN documentos d ON d.id = cl.documento_id
            JOIN empresas e ON e.id = d.empresa_id
            WHERE {' AND '.join(filtros)}
            ORDER BY e.nombre, cl.orden
            """,
            params,
        )
        filas = cur.fetchall()

    # Agrupado por empresa: una sola fila de comparador.php por empresa, con sus clausulas
    # aprobadas para este titulo (normalmente una, pero no se asume).
    por_empresa: dict = {}
    for fila in filas:
        emp = por_empresa.setdefault(
            fila["empresa_id"], {"empresa_id": fila["empresa_id"], "empresa_nombre": fila["empresa_nombre"], "clausulas": []}
        )
        emp["clausulas"].append({"id": fila["clausula_id"], "texto": fila["texto"], "documento_id": fila["documento_id"]})
    return list(por_empresa.values())


# ---------------------------------------------------------------------------
# Negociacion colectiva pre-firma (Art IV bis, Fase 3 / spec-negociacion.md) -- peticion
# (sindicato), oferta (empresa), reunion y acuerdo por titulo. Antecede al Art IV: una
# Empresa puede tener documentos que nunca pasaron por aca (convenciones cargadas directo).
# Ver la negociacion es AdminTenant/Revisor/Editor -- no es un reporte publicado (Art IV.9),
# asi que Visualizador no entra (spec-negociacion.md §4).
# ---------------------------------------------------------------------------

_VER_NEGOCIACION = ("AdminTenant", "Revisor", "Editor")
_EDITAR_NEGOCIACION = ("AdminTenant", "Editor")


def _registrar_evento_negociacion(cur, negociacion_id, evento: str, usuario_id, detalle: str = None) -> None:
    cur.execute(
        "INSERT INTO bitacora_negociacion (negociacion_id, evento, usuario_id, detalle) VALUES (%s, %s, %s, %s)",
        (negociacion_id, evento, usuario_id, detalle),
    )


def _obtener_negociacion(negociacion_id, tenant_id) -> dict:
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            """
            SELECT n.id, n.estado, n.fecha_inicio, n.fecha_cierre, n.empresa_id, e.nombre AS empresa_nombre
            FROM negociaciones n
            JOIN empresas e ON e.id = n.empresa_id
            WHERE n.id = %s AND n.tenant_id = %s
            """,
            (negociacion_id, tenant_id),
        )
        negociacion = cur.fetchone()
        if negociacion is None:
            raise HTTPException(404, "negociación no encontrada")

        cur.execute(
            """
            SELECT p.id, p.nro_peticion, p.texto, p.titulo_id, t.nombre AS titulo_nombre, p.created_at
            FROM peticiones p
            LEFT JOIN taxonomia_titulos t ON t.id = p.titulo_id
            WHERE p.negociacion_id = %s
            ORDER BY p.nro_peticion
            """,
            (negociacion_id,),
        )
        peticiones = cur.fetchall()
        if peticiones:
            cur.execute(
                "SELECT id, peticion_id, texto, created_at FROM ofertas WHERE peticion_id = ANY(%s) ORDER BY created_at",
                ([p["id"] for p in peticiones],),
            )
            ofertas_por_peticion: dict = {}
            for oferta in cur.fetchall():
                ofertas_por_peticion.setdefault(oferta["peticion_id"], []).append(oferta)
            for peticion in peticiones:
                peticion["ofertas"] = ofertas_por_peticion.get(peticion["id"], [])
        negociacion["peticiones"] = peticiones

        cur.execute(
            "SELECT id, fecha, asistentes, resumen, created_at FROM reuniones WHERE negociacion_id = %s ORDER BY fecha",
            (negociacion_id,),
        )
        negociacion["reuniones"] = cur.fetchall()

        cur.execute(
            """
            SELECT a.id, a.titulo_id, t.nombre AS titulo_nombre, a.texto_acordado,
                   a.peticion_id, a.oferta_id, a.created_at
            FROM acuerdos a
            JOIN taxonomia_titulos t ON t.id = a.titulo_id
            WHERE a.negociacion_id = %s
            ORDER BY a.created_at DESC
            """,
            (negociacion_id,),
        )
        negociacion["acuerdos"] = cur.fetchall()

        cur.execute(
            """
            SELECT id, estado, estado_detalle, version_negociacion, created_at
            FROM documentos WHERE negociacion_id = %s ORDER BY version_negociacion
            """,
            (negociacion_id,),
        )
        negociacion["documentos"] = cur.fetchall()

    return negociacion


@app.post("/negociaciones", status_code=201)
def crear_negociacion(request: Request, empresa_id: uuid.UUID = Form(...)):
    claims = require_role(request, *_EDITAR_NEGOCIACION)
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute("SELECT id FROM empresas WHERE id = %s AND tenant_id = %s", (empresa_id, claims.tenant_id))
        if cur.fetchone() is None:
            raise HTTPException(404, f"empresa_id {empresa_id} no existe en tu catálogo")

        cur.execute(
            "INSERT INTO negociaciones (tenant_id, empresa_id) VALUES (%s, %s) RETURNING id",
            (claims.tenant_id, empresa_id),
        )
        negociacion_id = cur.fetchone()["id"]
        _registrar_evento_negociacion(cur, negociacion_id, "creacion", claims.user_id)
        conn.commit()
    return _obtener_negociacion(negociacion_id, claims.tenant_id)


@app.get("/negociaciones")
def listar_negociaciones(request: Request, empresa_id: uuid.UUID):
    claims = require_role(request, *_VER_NEGOCIACION)
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            """
            SELECT n.id, n.estado, n.fecha_inicio, n.fecha_cierre, n.empresa_id, e.nombre AS empresa_nombre
            FROM negociaciones n
            JOIN empresas e ON e.id = n.empresa_id
            WHERE n.tenant_id = %s AND n.empresa_id = %s
            ORDER BY n.fecha_inicio DESC
            """,
            (claims.tenant_id, empresa_id),
        )
        return cur.fetchall()


@app.get("/negociaciones/{negociacion_id}")
def obtener_negociacion(request: Request, negociacion_id: uuid.UUID):
    claims = require_role(request, *_VER_NEGOCIACION)
    return _obtener_negociacion(negociacion_id, claims.tenant_id)


def _cargar_negociacion_abierta(cur, negociacion_id, tenant_id) -> dict:
    cur.execute(
        "SELECT id, estado FROM negociaciones WHERE id = %s AND tenant_id = %s",
        (negociacion_id, tenant_id),
    )
    negociacion = cur.fetchone()
    if negociacion is None:
        raise HTTPException(404, "negociación no encontrada")
    if negociacion["estado"] != "abierta":
        raise HTTPException(422, "la negociación está cerrada -- reabrila antes de agregar algo nuevo")
    return negociacion


@app.post("/negociaciones/{negociacion_id}/peticiones", status_code=201)
def crear_peticion(
    request: Request,
    negociacion_id: uuid.UUID,
    nro_peticion: int = Form(...),
    texto: str = Form(...),
    titulo_id: Optional[int] = Form(None),
):
    claims = require_role(request, *_EDITAR_NEGOCIACION)
    with get_conn() as conn, conn.cursor() as cur:
        _cargar_negociacion_abierta(cur, negociacion_id, claims.tenant_id)
        cur.execute(
            "INSERT INTO peticiones (negociacion_id, titulo_id, nro_peticion, texto) VALUES (%s, %s, %s, %s) RETURNING id",
            (negociacion_id, titulo_id, nro_peticion, texto),
        )
        peticion_id = cur.fetchone()["id"]
        _registrar_evento_negociacion(cur, negociacion_id, "peticion", claims.user_id, f"petición #{nro_peticion}")
        conn.commit()
    return {"id": peticion_id}


@app.post("/peticiones/{peticion_id}/ofertas", status_code=201)
def crear_oferta(request: Request, peticion_id: int, texto: str = Form(...)):
    claims = require_role(request, *_EDITAR_NEGOCIACION)
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            """
            SELECT p.negociacion_id FROM peticiones p
            JOIN negociaciones n ON n.id = p.negociacion_id
            WHERE p.id = %s AND n.tenant_id = %s
            """,
            (peticion_id, claims.tenant_id),
        )
        fila = cur.fetchone()
        if fila is None:
            raise HTTPException(404, "petición no encontrada")
        negociacion_id = fila["negociacion_id"]
        _cargar_negociacion_abierta(cur, negociacion_id, claims.tenant_id)

        cur.execute("INSERT INTO ofertas (peticion_id, texto) VALUES (%s, %s) RETURNING id", (peticion_id, texto))
        oferta_id = cur.fetchone()["id"]
        _registrar_evento_negociacion(cur, negociacion_id, "oferta", claims.user_id)
        conn.commit()
    return {"id": oferta_id}


@app.post("/negociaciones/{negociacion_id}/reuniones", status_code=201)
def crear_reunion(
    request: Request,
    negociacion_id: uuid.UUID,
    fecha: str = Form(...),
    asistentes: Optional[str] = Form(None),
    resumen: Optional[str] = Form(None),
):
    claims = require_role(request, *_EDITAR_NEGOCIACION)
    with get_conn() as conn, conn.cursor() as cur:
        _cargar_negociacion_abierta(cur, negociacion_id, claims.tenant_id)
        cur.execute(
            "INSERT INTO reuniones (negociacion_id, fecha, asistentes, resumen) VALUES (%s, %s, %s, %s) RETURNING id",
            (negociacion_id, fecha, asistentes, resumen),
        )
        reunion_id = cur.fetchone()["id"]
        _registrar_evento_negociacion(cur, negociacion_id, "reunion", claims.user_id)
        conn.commit()
    return {"id": reunion_id}


@app.post("/negociaciones/{negociacion_id}/acuerdos", status_code=201)
def crear_acuerdo(
    request: Request,
    negociacion_id: uuid.UUID,
    titulo_id: int = Form(...),
    texto_acordado: str = Form(...),
    peticion_id: Optional[int] = Form(None),
    oferta_id: Optional[int] = Form(None),
):
    claims = require_role(request, *_EDITAR_NEGOCIACION)
    with get_conn() as conn, conn.cursor() as cur:
        _cargar_negociacion_abierta(cur, negociacion_id, claims.tenant_id)
        cur.execute("SELECT id FROM taxonomia_titulos WHERE id = %s", (titulo_id,))
        if cur.fetchone() is None:
            raise HTTPException(422, f"titulo_id {titulo_id} no existe en la taxonomía")

        # peticion_id/oferta_id son opcionales, pero si se mandan deben pertenecer a ESTA
        # negociacion -- si no, cualquier id ajeno (incluso de otro tenant) quedaria
        # colgado de un acuerdo como referencia decorativa (Art VI.2).
        if peticion_id is not None:
            cur.execute(
                "SELECT id FROM peticiones WHERE id = %s AND negociacion_id = %s",
                (peticion_id, negociacion_id),
            )
            if cur.fetchone() is None:
                raise HTTPException(422, f"peticion_id {peticion_id} no pertenece a esta negociación")
        if oferta_id is not None:
            cur.execute(
                "SELECT o.id FROM ofertas o JOIN peticiones p ON p.id = o.peticion_id WHERE o.id = %s AND p.negociacion_id = %s",
                (oferta_id, negociacion_id),
            )
            if cur.fetchone() is None:
                raise HTTPException(422, f"oferta_id {oferta_id} no pertenece a esta negociación")

        cur.execute(
            """
            INSERT INTO acuerdos (negociacion_id, titulo_id, texto_acordado, peticion_id, oferta_id)
            VALUES (%s, %s, %s, %s, %s) RETURNING id
            """,
            (negociacion_id, titulo_id, texto_acordado, peticion_id, oferta_id),
        )
        acuerdo_id = cur.fetchone()["id"]
        _registrar_evento_negociacion(cur, negociacion_id, "acuerdo", claims.user_id)
        conn.commit()
    return {"id": acuerdo_id}


def _armar_docx_acuerdos(acuerdos: list[dict]) -> bytes:
    """Un documento sintetico, un titulo por clausula, a partir del acuerdo mas reciente de
    cada titulo (spec-negociacion.md §5) -- se persiste y procesa igual que un archivo
    cargado, sin rama especial en el pipeline (extraction.py ya sabe leer .docx).

    El prefijo "CLAUSULA -- <titulo>" no es cosmetico: app/segmentation.py detecta limites
    de clausula buscando el patron "CLAUSULA"/"ARTICULO" (no encabezados de Word, que
    extract_text descarta -- doc.paragraphs los aplana igual que cualquier parrafo). Sin
    ese prefijo, dos o mas acuerdos terminaban fusionados en una sola clausula porque
    extract_docx los une con un solo salto de linea (sin linea en blanco de por medio, el
    fallback de parrafos tampoco los separaba)."""
    import io

    from docx import Document as DocxDocument

    doc = DocxDocument()
    for acuerdo in acuerdos:
        doc.add_paragraph(f"CLAUSULA -- {acuerdo['titulo_nombre']}")
        doc.add_paragraph(acuerdo["texto_acordado"])
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@app.post("/negociaciones/{negociacion_id}/cerrar", status_code=201)
def cerrar_negociacion(request: Request, negociacion_id: uuid.UUID, background: BackgroundTasks):
    claims = require_role(request, "AdminTenant")
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            "SELECT id, estado, empresa_id FROM negociaciones WHERE id = %s AND tenant_id = %s",
            (negociacion_id, claims.tenant_id),
        )
        negociacion = cur.fetchone()
        if negociacion is None:
            raise HTTPException(404, "negociación no encontrada")
        if negociacion["estado"] != "abierta":
            raise HTTPException(422, "la negociación ya está cerrada")

        # El acuerdo mas reciente por titulo dentro de esta negociacion (DISTINCT ON,
        # ordenado por created_at DESC) -- asi un addendum que renegocia un solo titulo no
        # duplica los demas al re-cerrar (spec-negociacion.md §5).
        cur.execute(
            """
            SELECT DISTINCT ON (a.titulo_id) a.titulo_id, t.nombre AS titulo_nombre, a.texto_acordado
            FROM acuerdos a
            JOIN taxonomia_titulos t ON t.id = a.titulo_id
            WHERE a.negociacion_id = %s
            ORDER BY a.titulo_id, a.created_at DESC
            """,
            (negociacion_id,),
        )
        acuerdos_vigentes = cur.fetchall()
        if not acuerdos_vigentes:
            raise HTTPException(422, "no hay acuerdos registrados -- no se puede cerrar sin al menos uno")

        cur.execute("SELECT COUNT(*) AS n FROM documentos WHERE negociacion_id = %s", (negociacion_id,))
        version = cur.fetchone()["n"] + 1

        contenido = _armar_docx_acuerdos(acuerdos_vigentes)
        nombre_archivo = f"negociacion_{negociacion_id}_v{version}.docx"
        ruta_archivo = storage.guardar(claims.tenant_id, nombre_archivo, contenido)

        cur.execute(
            """
            INSERT INTO documentos (tenant_id, empresa_id, origen, ruta_archivo, estado, negociacion_id, version_negociacion)
            VALUES (%s, %s, 'negociacion', %s, 'pendiente', %s, %s)
            RETURNING id
            """,
            (claims.tenant_id, negociacion["empresa_id"], ruta_archivo, negociacion_id, version),
        )
        documento_id = cur.fetchone()["id"]

        cur.execute(
            "UPDATE negociaciones SET estado = 'cerrada', fecha_cierre = now() WHERE id = %s",
            (negociacion_id,),
        )
        _registrar_evento_negociacion(cur, negociacion_id, "cierre", claims.user_id, f"documento v{version}")
        conn.commit()

    background.add_task(_procesar_pipeline, documento_id, claims.tenant_id, contenido, ".docx")
    return _obtener_negociacion(negociacion_id, claims.tenant_id)


@app.post("/negociaciones/{negociacion_id}/reabrir")
def reabrir_negociacion(request: Request, negociacion_id: uuid.UUID):
    claims = require_role(request, "AdminTenant")
    with get_conn() as conn, conn.cursor() as cur:
        cur.execute(
            "UPDATE negociaciones SET estado = 'abierta', fecha_cierre = NULL WHERE id = %s AND tenant_id = %s AND estado = 'cerrada' RETURNING id",
            (negociacion_id, claims.tenant_id),
        )
        if cur.fetchone() is None:
            raise HTTPException(404, "negociación no encontrada, o ya está abierta")
        _registrar_evento_negociacion(cur, negociacion_id, "reapertura", claims.user_id)
        conn.commit()
    return _obtener_negociacion(negociacion_id, claims.tenant_id)
