import io

import pymupdf as fitz
import pytest
from docx import Document as DocxDocument

from app.extraction import ExtractionError, extract_text


def test_extension_no_soportada_lanza_extraction_error():
    with pytest.raises(ExtractionError, match="no soportado"):
        extract_text(b"no deberia procesarse", ".txt")


def test_extrae_texto_nativo_de_docx():
    doc = DocxDocument()
    doc.add_paragraph("CLAUSULA PRIMERA: Objeto del contrato.")
    doc.add_paragraph("")  # parrafo vacio — no debe aparecer en el resultado
    doc.add_paragraph("CLAUSULA SEGUNDA: Vigencia.")
    buffer = io.BytesIO()
    doc.save(buffer)

    texto = extract_text(buffer.getvalue(), ".docx")

    assert "CLAUSULA PRIMERA: Objeto del contrato." in texto
    assert "CLAUSULA SEGUNDA: Vigencia." in texto
    # Sin lineas en blanco colgando de parrafos vacios (Document.paragraphs los filtra).
    assert "\n\n" not in texto


def test_extrae_texto_nativo_de_pdf():
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "CLAUSULA PRIMERA texto nativo suficiente para no disparar OCR " * 3)
    contenido = doc.tobytes()
    doc.close()

    texto = extract_text(contenido, ".pdf")

    assert "CLAUSULA PRIMERA" in texto
